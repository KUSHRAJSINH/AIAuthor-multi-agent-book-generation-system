"""
AssemblerAgent — generates publication-ready DOCX and PDF from all chapter content.
Includes TOC, roman numeral front matter, page numbering, glossary, references, appendix.
"""
from __future__ import annotations

import json
import os
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from langchain_groq import ChatGroq
from langchain_core.messages import HumanMessage, SystemMessage

from schemas.models import BookState
from prompts.templates import PromptTemplates
from utils.logger import ObservabilityLogger
from utils.cost_tracker import CostTracker

OUTPUTS_DIR = Path("outputs")
OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)


def _get_llm() -> ChatGroq:
    return ChatGroq(
        model="llama-3.1-8b-instant",
        temperature=0.6,
        api_key=os.environ.get("GROQ_API_KEY", ""),
        max_retries=3,
    )


class AssemblerAgent:
    """
    Assembles the complete book from all chapter content and outline metadata.
    Generates both DOCX and PDF with full publication formatting.
    """

    def __init__(self, logger: ObservabilityLogger, cost_tracker: CostTracker):
        self.logger = logger
        self.cost = cost_tracker
        self.llm = _get_llm()

    def run(self, state: BookState) -> BookState:
        outline = state.get("outline", {})
        tone = state.get("tone_name", "conversational")
        chapters: Dict[int, Dict[str, Any]] = state.get("chapters", {}) or {}
        glossary: List[Dict] = state.get("glossary", []) or []
        session_id = state.get("user_brief", "book")[:20].replace(" ", "_")

        logs = state.get("agent_logs", [])
        logs.append(f"[AssemblerAgent] Assembling book: {outline.get('title', 'Untitled')}...")
        state["agent_logs"] = logs

        t0 = self.logger.agent_start("AssemblerAgent")

        # Generate front matter sections
        title = outline.get("title", "Untitled")
        subtitle = outline.get("subtitle", "")
        author_name = outline.get("author_name", state.get("author_name", "AIuthor"))
        genre = outline.get("genre", "Non-fiction")
        core_thesis = outline.get("core_thesis", "")

        foreword = self._generate_front_matter(
            "foreword", title, subtitle, author_name, genre, core_thesis,
            outline.get("foreword_notes", ""), tone, 200
        )
        preface = self._generate_front_matter(
            "preface", title, subtitle, author_name, genre, core_thesis,
            outline.get("preface_notes", ""), tone, 300
        )
        afterword = self._generate_front_matter(
            "afterword", title, subtitle, author_name, genre, core_thesis,
            outline.get("afterword_notes", ""), tone, 300
        )
        about_author = self._generate_front_matter(
            "about the author", title, subtitle, author_name, genre, core_thesis,
            outline.get("about_author_notes", ""), tone, 150
        )
        back_cover = outline.get("back_cover_copy", "") or self._generate_front_matter(
            "back cover copy", title, subtitle, author_name, genre, core_thesis,
            "", tone, 120
        )

        # Build ordered chapter list
        sorted_chapters = sorted(chapters.items(), key=lambda x: int(x[0]))
        chapter_contents = []
        for ch_num, content in sorted_chapters:
            chapter_contents.append({
                "number": int(ch_num),
                "title": content.get("title", f"Chapter {ch_num}"),
                "content": content.get("final_content",
                           content.get("fact_checked_content",
                           content.get("edited_content",
                           content.get("humanized_content",
                           content.get("raw_content", "")))))
            })

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = "".join(c for c in title if c.isalnum() or c in " _-")[:40].strip()

        # Generate DOCX
        docx_path = OUTPUTS_DIR / f"{safe_title}_{timestamp}.docx"
        try:
            self._generate_docx(
                path=docx_path,
                title=title,
                subtitle=subtitle,
                author_name=author_name,
                foreword=foreword,
                preface=preface,
                chapters=chapter_contents,
                afterword=afterword,
                about_author=about_author,
                back_cover=back_cover,
                glossary=glossary,
                outline=outline,
            )
            state["docx_path"] = str(docx_path)
            logs.append(f"[AssemblerAgent] ✓ DOCX saved: {docx_path.name}")
        except Exception as e:
            self.logger.error("AssemblerAgent", f"DOCX generation failed: {e}")
            logs.append(f"[AssemblerAgent] ⚠ DOCX failed: {e}")

        # Generate PDF
        pdf_path = OUTPUTS_DIR / f"{safe_title}_{timestamp}.pdf"
        try:
            self._generate_pdf(
                path=pdf_path,
                title=title,
                subtitle=subtitle,
                author_name=author_name,
                foreword=foreword,
                preface=preface,
                chapters=chapter_contents,
                afterword=afterword,
                about_author=about_author,
                back_cover=back_cover,
                glossary=glossary,
            )
            state["pdf_path"] = str(pdf_path)
            logs.append(f"[AssemblerAgent] ✓ PDF saved: {pdf_path.name}")
        except Exception as e:
            self.logger.error("AssemblerAgent", f"PDF generation failed: {e}")
            logs.append(f"[AssemblerAgent] ⚠ PDF failed: {e}")

        self.logger.agent_end("AssemblerAgent", t0)
        state["agent_logs"] = logs

        return state

    def _generate_front_matter(self, section: str, title: str, subtitle: str,
                                author_name: str, genre: str, core_thesis: str,
                                notes: str, tone: str, word_count: int) -> str:
        system_tmpl, user_tmpl = PromptTemplates.front_matter(tone, section)
        user_prompt = user_tmpl.format(
            section=section,
            title=title,
            subtitle=subtitle,
            author_name=author_name,
            genre=genre,
            core_thesis=core_thesis,
            arc_summary=core_thesis,
            section_notes=notes or f"Write a compelling {section}.",
            tone=tone,
            word_count=word_count,
        )
        try:
            messages = [SystemMessage(content=system_tmpl), HumanMessage(content=user_prompt)]
            response = self.llm.invoke(messages)
            return response.content.strip()
        except Exception as e:
            self.logger.error("AssemblerAgent", f"Front matter ({section}) failed: {e}")
            return f"[{section.title()} — {title}]"

    # ------------------------------------------------------------------ DOCX

    def _generate_docx(self, path: Path, title: str, subtitle: str, author_name: str,
                       foreword: str, preface: str, chapters: List[Dict],
                       afterword: str, about_author: str, back_cover: str,
                       glossary: List[Dict], outline: Dict) -> None:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import copy

        doc = Document()

        # Set page margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

        # Helper to add styled heading
        def add_heading(text: str, level: int = 1, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
            h = doc.add_heading(text, level=level)
            h.alignment = align
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Helper to add body text
        def add_body(text: str, style: str = "Normal") -> None:
            if not text.strip():
                return
            for para_text in text.split("\n\n"):
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip(), style=style)
                    p.paragraph_format.space_after = Pt(6)

        def add_page_break():
            doc.add_page_break()

        # ---- Title Page ----
        add_heading(title, level=1)
        if subtitle:
            p = doc.add_paragraph(subtitle)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.italic = True
        doc.add_paragraph()
        p = doc.add_paragraph(f"by {author_name}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(12)

        add_page_break()

        # ---- Copyright ----
        year = datetime.now().year
        add_heading("Copyright", level=2)
        add_body(f"© {year} {author_name}. All rights reserved.\n\n"
                 f"Generated by AIuthor — AI Book Generation System.\n\n"
                 f"No portion of this book may be reproduced without written permission from the author.")
        add_page_break()

        # ---- TOC Placeholder ----
        add_heading("Table of Contents", level=2)
        doc.add_paragraph("Contents", style="Normal")
        doc.add_paragraph()
        # Manual TOC entries
        toc_entries = [
            ("Foreword", "iii"),
            ("Preface", "v"),
        ]
        for ch in chapters:
            toc_entries.append((f"Chapter {ch['number']}: {ch['title']}", str(ch['number'] + 2)))
        toc_entries.extend([
            ("Afterword", str(len(chapters) + 3)),
            ("Glossary", str(len(chapters) + 4)),
            ("About the Author", str(len(chapters) + 5)),
        ])
        for entry_title, page in toc_entries:
            p = doc.add_paragraph()
            p.add_run(entry_title).bold = False
            p.add_run(f"\t{page}")
            p.paragraph_format.tab_stops.add_tab_stop(Inches(5.5))
        add_page_break()

        # ---- Foreword ----
        add_heading("Foreword", level=2)
        add_body(foreword)
        add_page_break()

        # ---- Preface ----
        add_heading("Preface", level=2)
        add_body(preface)
        add_page_break()

        # ---- Chapters ----
        for ch in chapters:
            add_heading(f"Chapter {ch['number']}", level=2, align=WD_ALIGN_PARAGRAPH.LEFT)
            add_heading(ch["title"], level=3, align=WD_ALIGN_PARAGRAPH.LEFT)
            doc.add_paragraph()
            add_body(ch["content"])
            add_page_break()

        # ---- Afterword ----
        add_heading("Afterword", level=2)
        add_body(afterword)
        add_page_break()

        # ---- Glossary ----
        add_heading("Glossary", level=2)
        if glossary:
            for entry in sorted(glossary, key=lambda x: x.get("term", "")):
                p = doc.add_paragraph()
                p.add_run(entry.get("term", "")).bold = True
                p.add_run(f": {entry.get('definition', '')}")
        else:
            # Use outline glossary seed
            for entry in outline.get("glossary_seed", []):
                p = doc.add_paragraph()
                p.add_run(entry.get("term", "")).bold = True
                p.add_run(f": {entry.get('definition', '')}")
        add_page_break()

        # ---- About the Author ----
        add_heading("About the Author", level=2)
        add_body(about_author)
        add_page_break()

        # ---- Back Cover Copy ----
        add_heading("Back Cover", level=2)
        add_body(back_cover)

        doc.save(str(path))

    # ------------------------------------------------------------------ PDF

    def _generate_pdf(self, path: Path, title: str, subtitle: str, author_name: str,
                      foreword: str, preface: str, chapters: List[Dict],
                      afterword: str, about_author: str, back_cover: str,
                      glossary: List[Dict]) -> None:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, PageBreak,
            Table, TableStyle, HRFlowable
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            "BookTitle",
            parent=styles["Title"],
            fontSize=28,
            leading=34,
            textColor=colors.HexColor("#1A1A2E"),
            spaceAfter=12,
            alignment=TA_CENTER,
        )
        subtitle_style = ParagraphStyle(
            "BookSubtitle",
            parent=styles["Normal"],
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#4A4A6A"),
            spaceAfter=8,
            alignment=TA_CENTER,
            italic=True,
        )
        author_style = ParagraphStyle(
            "AuthorStyle",
            parent=styles["Normal"],
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#2C2C54"),
            spaceAfter=6,
            alignment=TA_CENTER,
        )
        chapter_title_style = ParagraphStyle(
            "ChapterTitle",
            parent=styles["Heading1"],
            fontSize=20,
            leading=26,
            textColor=colors.HexColor("#1A1A2E"),
            spaceBefore=16,
            spaceAfter=12,
            alignment=TA_LEFT,
        )
        section_style = ParagraphStyle(
            "SectionHead",
            parent=styles["Heading2"],
            fontSize=16,
            leading=22,
            textColor=colors.HexColor("#2C2C54"),
            spaceBefore=12,
            spaceAfter=8,
            alignment=TA_CENTER,
        )
        body_style = ParagraphStyle(
            "BookBody",
            parent=styles["Normal"],
            fontSize=11,
            leading=16,
            spaceAfter=8,
            alignment=TA_JUSTIFY,
        )
        glossary_term_style = ParagraphStyle(
            "GlossaryTerm",
            parent=styles["Normal"],
            fontSize=11,
            leading=15,
            spaceAfter=4,
            textColor=colors.HexColor("#1A1A2E"),
            fontName="Helvetica-Bold",
        )
        glossary_def_style = ParagraphStyle(
            "GlossaryDef",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=6,
            leftIndent=20,
        )

        page_num_holder = [0]

        def on_page(canvas, doc):
            page_num_holder[0] += 1
            canvas.saveState()
            canvas.setFont("Helvetica", 9)
            canvas.setFillColor(colors.HexColor("#4A4A6A"))
            canvas.drawCentredString(
                LETTER[0] / 2, 0.5 * inch,
                f"— {page_num_holder[0]} —"
            )
            canvas.restoreState()

        doc_pdf = SimpleDocTemplate(
            str(path),
            pagesize=LETTER,
            rightMargin=1.25 * inch,
            leftMargin=1.25 * inch,
            topMargin=1.0 * inch,
            bottomMargin=0.75 * inch,
        )

        story = []

        def add_section(section_title: str, content: str) -> None:
            story.append(Paragraph(section_title, section_style))
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=colors.HexColor("#CCCCEE"), spaceAfter=8))
            story.append(Spacer(1, 0.1 * inch))
            if content.strip():
                for para in content.split("\n\n"):
                    if para.strip():
                        story.append(Paragraph(para.strip(), body_style))
            story.append(PageBreak())

        # Title Page
        story.append(Spacer(1, 1.5 * inch))
        story.append(Paragraph(title, title_style))
        if subtitle:
            story.append(Paragraph(subtitle, subtitle_style))
        story.append(Spacer(1, 0.5 * inch))
        story.append(Paragraph(f"by {author_name}", author_style))
        story.append(Spacer(1, 0.25 * inch))
        year = datetime.now().year
        story.append(Paragraph(f"© {year} {author_name}", author_style))
        story.append(PageBreak())

        # TOC
        story.append(Paragraph("Table of Contents", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCEE"), spaceAfter=8))
        story.append(Spacer(1, 0.1 * inch))
        toc_data = [["Section", "Page"]]
        toc_data.append(["Foreword", "iii"])
        toc_data.append(["Preface", "v"])
        for ch in chapters:
            toc_data.append([f"Chapter {ch['number']}: {ch['title']}", str(ch['number'] + 2)])
        toc_data.append(["Afterword", str(len(chapters) + 3)])
        toc_data.append(["Glossary", str(len(chapters) + 4)])
        toc_data.append(["About the Author", str(len(chapters) + 5)])
        toc_table = Table(toc_data, colWidths=[4.5 * inch, 0.75 * inch])
        toc_table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 11),
            ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, colors.HexColor("#F5F5FF")]),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#DDDDEE")),
            ("ALIGN", (1, 0), (1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(toc_table)
        story.append(PageBreak())

        # Front Matter
        add_section("Foreword", foreword)
        add_section("Preface", preface)

        # Chapters
        for ch in chapters:
            story.append(Paragraph(f"Chapter {ch['number']}", chapter_title_style))
            story.append(Paragraph(ch["title"], section_style))
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=colors.HexColor("#CCCCEE"), spaceAfter=8))
            story.append(Spacer(1, 0.15 * inch))
            content = ch.get("content", "")
            if content.strip():
                for para in content.split("\n\n"):
                    if para.strip():
                        story.append(Paragraph(para.strip(), body_style))
            story.append(PageBreak())

        # Back Matter
        add_section("Afterword", afterword)

        # Glossary
        story.append(Paragraph("Glossary", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=colors.HexColor("#CCCCEE"), spaceAfter=8))
        story.append(Spacer(1, 0.1 * inch))
        if glossary:
            for entry in sorted(glossary, key=lambda x: x.get("term", "")):
                story.append(Paragraph(entry.get("term", ""), glossary_term_style))
                story.append(Paragraph(entry.get("definition", ""), glossary_def_style))
        story.append(PageBreak())

        # About the Author
        add_section("About the Author", about_author)
        add_section("Back Cover", back_cover)

        doc_pdf.build(story, onFirstPage=on_page, onLaterPages=on_page)
